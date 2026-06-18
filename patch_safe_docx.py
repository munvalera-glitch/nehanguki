import glob
from docx import Document
from docx.shared import Pt

for filepath in glob.glob("templates/schedule_*.docx"):
    doc = Document(filepath)
    for p in doc.paragraphs:
        if "___ 명의 일본 체류일정표는 다음과 같습니다." in p.text:
            p.text = p.text.replace("___", "{companionCount}")
            
        if "Name of companion and relationship with applicant" in p.text:
            # We want to replace the companions list
            pass
            
        if "1. ___________________  (                         )  2. ___________________  (                         )" in p.text:
            p.clear()
            run = p.add_run("{#companions}\n{index}. {name}\n{/companions}")
            run.font.name = "Arial"
            run.font.size = Pt(10)
            
        if "3. ___________________  (                         )  4. ___________________  (                         )" in p.text:
            p.text = ""
        if "5. ___________________  (                         )" in p.text:
            p.text = ""

    doc.save(filepath)
    print(f"Patched {filepath}")
