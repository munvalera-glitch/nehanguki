import glob
from docx import Document
import re

def replace_text_in_paragraph(paragraph, old_text, new_text):
    # docx splits text into runs which makes naive string replacement hard.
    # A simple trick is to concatenate all text, check if old_text is there,
    # and if so, clear all runs except the first one and set its text.
    # This might lose some formatting within the paragraph if it's mixed, 
    # but for simple placeholders like 'Day 1' it's fine.
    
    # Actually, simpler: replace if the old_text is exactly within a run
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
            
    # If not found in a single run, try full text replacement on the first run
    if old_text in paragraph.text:
        text = paragraph.text
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                run.text = text.replace(old_text, new_text)
            else:
                run.text = ""
        return True
    return False

def process_file(filepath):
    doc = Document(filepath)
    
    # Update Paragraphs (for Date and Applicant Name)
    for p in doc.paragraphs:
        if "2026 (Year) (Month) (Day)" in p.text:
            replace_text_in_paragraph(p, "2026 (Year) (Month) (Day)", "{currentDate}")
            
        if "신청인인" in p.text and "외" in p.text:
            # We want to replace "신청인인 _______________ 외" with "신청인인 {applicantName} 외"
            text = p.text
            # Use regex to find "신청인인 _+ 외"
            new_text = re.sub(r'신청인인\s*_+.*?외', '신청인인 {applicantName} 외', text)
            if new_text != text:
                for i, run in enumerate(p.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ""
                        
    # Update Tables (for Day 1, Day 2...)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for i in range(1, 6):
                        day_str = f"Day {i}"
                        if day_str in p.text:
                            replace_text_in_paragraph(p, day_str, f"{{date{i}}}")

    doc.save(filepath)
    print(f"Updated {filepath}")

files = glob.glob("templates/schedule_*.docx")
for f in files:
    try:
        process_file(f)
    except Exception as e:
        print(f"Error processing {f}: {e}")
