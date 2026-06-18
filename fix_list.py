import glob
from docx import Document
import re

def process_file(filepath):
    doc = Document(filepath)
    modified = False
    
    # 1. Handle `___ 명의` (convert to `{companionCount} 명의`)
    # The text is: 신청인인 {applicantName} 외 ____ 명의
    for p in doc.paragraphs:
        if "명의" in p.text and "외" in p.text:
            text = p.text
            new_text = re.sub(r'외\s*_+.*?명의', '외 {companionCount} 명의', text)
            if new_text != text:
                for i, run in enumerate(p.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ""
                modified = True
                
    # 2. Handle the list 1. 2. 3. 4. 5.
    paragraphs_to_remove = []
    found_first = False
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("1.") and not found_first:
            found_first = True
            # Replace 1. with the loop
            for idx, run in enumerate(p.runs):
                if idx == 0:
                    run.text = "{#companions}\n{index}. {name} ({relationship})\n{/companions}"
                else:
                    run.text = ""
            modified = True
        elif text in ["2.", "3.", "4.", "5."] and found_first:
            paragraphs_to_remove.append(p)
            
    for p in paragraphs_to_remove:
        # Removing paragraph from docx
        p._element.getparent().remove(p._element)
        modified = True

    if modified:
        doc.save(filepath)
        print(f"Fixed list in {filepath}")

files = glob.glob("templates/schedule_*.docx")
for f in files:
    try:
        process_file(f)
    except Exception as e:
        print(f"Error processing {f}: {e}")
